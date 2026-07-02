from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def export_proposal_docx(state: dict, output_path: str) -> str:
    doc = Document()
    
    # Professional Styling
    # Configure styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Off-black
    
    # 1. Cover Page
    # Space at top
    for _ in range(5):
        doc.add_paragraph()
        
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("BID PROPOSAL DRAFT")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy
    
    doc.add_paragraph()
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(state.get("tender_title", "Government Tender Project"))
    sub_run.font.size = Pt(16)
    sub_run.font.italic = True
    
    for _ in range(8):
        doc.add_paragraph()
        
    profile = state.get("business_profile") or {}
    company_name = profile.get("company_name") or profile.get("name") or "TechVenture Solutions"
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run(
        f"Prepared by: {company_name}\n"
        f"Date: {datetime.now().strftime('%d %B %Y')}\n"
        f"Proposal ID: {state.get('analysis_id', 'ANL-001')}"
    )
    info_run.font.size = Pt(12)
    info_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    doc.add_page_break()
    
    # 2. Executive Summary
    h1 = doc.add_heading("1. Executive Summary", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    proposal = state.get("proposal_draft") or {}
    exec_summary = proposal.get("executive_summary") or "No executive summary generated."
    doc.add_paragraph(exec_summary)
    
    doc.add_paragraph()
    
    # 3. Technical Proposal
    h2 = doc.add_heading("2. Technical Proposal & Approach", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    tech_prop = proposal.get("technical_proposal") or "No technical proposal generated."
    doc.add_paragraph(tech_prop)
    
    doc.add_paragraph()
    
    # 4. Scope of Work
    h3 = doc.add_heading("3. Scope of Work", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    scope_work = proposal.get("scope_of_work") or "No scope of work generated."
    doc.add_paragraph(scope_work)
    
    doc.add_paragraph()
    
    # 5. Project Plan & Timeline
    h4 = doc.add_heading("4. Project Plan & Implementation Timeline", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    proj_plan = proposal.get("project_plan") or "No project plan generated."
    doc.add_paragraph(proj_plan)
    
    doc.add_paragraph()
    
    # 6. Budget & Cost Breakdown
    h5 = doc.add_heading("5. Financial Quote & Budget Template", level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    cost_est = state.get("cost_estimate") or {}
    budget_template = proposal.get("budget_template") or ""
    
    if budget_template:
        doc.add_paragraph(budget_template)
    else:
        # Generate a neat budget table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Item/Component'
        hdr_cells[1].text = 'Estimated Cost (INR)'
        
        items = [
            ('Development Cost', cost_est.get('development_cost', 0)),
            ('Infrastructure Cost', cost_est.get('infrastructure_cost', 0)),
            ('Team & Resource Cost', cost_est.get('team_cost', 0)),
            ('Operational & Maintenance Cost', cost_est.get('operational_cost', 0)),
            ('Total Project Bid Price', cost_est.get('total', 0))
        ]
        
        for name, cost in items:
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = f"₹ {cost:,.2f}"
            
    doc.add_paragraph()
    
    # 7. Compliance Checklist
    h6 = doc.add_heading("6. Compliance & Submission Checklist", level=1)
    h6.runs[0].font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    checklist = proposal.get("compliance_checklist") or []
    if checklist:
        doc.add_paragraph("The following list of documents and certifications must be checked and signed off prior to submission:")
        for item in checklist:
            p = doc.add_paragraph(style='List Bullet')
            # Add a checkbox symbol
            p.add_run("☐  ").font.bold = True
            p.add_run(item)
    else:
        doc.add_paragraph("No specific checklist items defined. Please ensure Standard company registrations, PAN, GST and audit sheets are ready.")
        
    doc.save(output_path)
    return output_path
